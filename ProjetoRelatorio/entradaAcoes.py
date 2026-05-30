#! python3
import datetime, os
from docx import Document
from docx.shared import Pt, Inches

# Criar caminho e abrir Relatório mensal recente
pasta = r"C:\Users\usuario\OneDrive - Universidade Federal de Pernambuco\FACULDADE\Relatórios PIBID 2024-26\2026"
dataHoje = datetime.datetime.now()
nomeArquivo = f"Relatório({dataHoje.strftime('%m.%Y')})_PIBID_LEONARDO_FERNANDO_DE_LEMOS_SCHULER.docx"
caminhoArquivo = os.path.join(pasta, nomeArquivo)

doc = Document(caminhoArquivo)

while True:
    print("Escolha sua entrada: \n1- ATIVIDADES PRINCIPAIS\n2- ATIVIDADES EXTRAS")
    escolha = int(input())
    if escolha == 1:
        print("ENTRADAS DE ATIVIDADES PRINCIPAIS")
        x = 3
    elif escolha == 2:
        print("ENTRADA DE ATIVIDADES EXTRAS")
        x = 5
    else: 
        print('Ops! Você deve escolher entre 1 e 2')
        continue
    # Acessar 4ª tabela do documento (seção 4.1 Atividades do PIBID realizadas)
    tabelaAtividades = doc.tables[x]

    # Criar linha nova para entrada
    novaLinha = tabelaAtividades.add_row()
    celulaUnica = novaLinha.cells[0]


    # Escrever primeiro parágrafo da entrada (negrito)
    tituloAtividade = input("Indique o título da atividade realizada: ")
    p1 = celulaUnica.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.line_spacing = 1.0
    run1 = p1.add_run(tituloAtividade)
    run1.bold = True

    # Escrever segundo parágrafo na mesma célula (recuo na 1º linha)
    descricaoAtividade = input("Agora descreva a atividade: ")
    p2 = celulaUnica.add_paragraph()
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    run2 = p2.add_run(descricaoAtividade)

    #Salvar arquivo
    doc.save(caminhoArquivo)
    print('Entrada salva com sucesso.')
    print('Gostaria de fazer uma nova entrada? \n1 - SIM\n2 - NÃO')
    continuar = input()
    if continuar != 1:
        break

finalizar = input()